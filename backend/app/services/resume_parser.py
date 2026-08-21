import io
import re
import logging
from typing import Tuple

logger = logging.getLogger("placemind.resume_parser")

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB

def clean_extracted_text(text: str) -> str:
    """Clean and normalize raw extracted resume text."""
    if not text:
        return ""
    # Replace multiple whitespaces/newlines with single space/newline
    text = re.sub(r'[\r\t\f\v]', ' ', text)
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\n+', '\n', text)
    return text.strip()

def parse_pdf_bytes(file_bytes: bytes) -> str:
    """Extract raw text from PDF binary bytes."""
    extracted_pages = []
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                extracted_pages.append(page_text)
    except Exception as e:
        logger.warning("pypdf extraction failed, attempting fallback text decoding: %s", str(e))
        # Fallback string extraction for simple PDFs
        raw_decoded = file_bytes.decode("latin-1", errors="ignore")
        text_matches = re.findall(r'\((.*?)\)\s*TJ|\((.*?)\)\s*Tj', raw_decoded)
        for match in text_matches:
            found = match[0] or match[1]
            if len(found) > 2:
                extracted_pages.append(found)

    full_text = "\n".join(extracted_pages)
    return clean_extracted_text(full_text)

def parse_docx_bytes(file_bytes: bytes) -> str:
    """Extract raw text from DOCX binary bytes."""
    extracted_paragraphs = []
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        for p in doc.paragraphs:
            if p.text:
                extracted_paragraphs.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        extracted_paragraphs.append(cell.text)
    except Exception as e:
        logger.warning("docx extraction failed, attempting fallback zip XML parsing: %s", str(e))
        import zipfile
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                xml_content = z.read("word/document.xml").decode("utf-8", errors="ignore")
                text_matches = re.findall(r'<w:t[^>]*>(.*?)</w:t>', xml_content)
                extracted_paragraphs.extend(text_matches)
        except Exception as ze:
            logger.error("ZIP XML parsing failed for docx: %s", str(ze))

    full_text = "\n".join(extracted_paragraphs)
    return clean_extracted_text(full_text)

def parse_resume_document(file_bytes: bytes, filename: str, content_type: str = "") -> Tuple[str, str]:
    """
    Validates file size and type, then parses text content.
    Returns (extracted_text, detected_file_type).
    Raises ValueError for unsupported formats, oversized files, or empty extraction.
    """
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise ValueError("File size exceeds maximum allowed limit of 10 MB.")

    if len(file_bytes) == 0:
        raise ValueError("Uploaded file is empty.")

    lower_name = filename.lower()
    is_pdf = lower_name.endswith(".pdf") or "pdf" in content_type.lower()
    is_docx = lower_name.endswith(".docx") or "word" in content_type.lower() or "officedocument" in content_type.lower()

    if not is_pdf and not is_docx:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")

    extracted_text = ""
    file_type = "pdf" if is_pdf else "docx"

    if is_pdf:
        extracted_text = parse_pdf_bytes(file_bytes)
    elif is_docx:
        extracted_text = parse_docx_bytes(file_bytes)

    if not extracted_text or len(extracted_text.strip()) < 10:
        raise ValueError("Failed to extract readable text from resume. Ensure the file contains text and is not encrypted.")

    return extracted_text, file_type
